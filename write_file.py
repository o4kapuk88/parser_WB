from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime
import locale


def replace_placeholder_in_doc(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    # Сохраняем стиль текста
                    style = run.style
                    # Заменяем текст
                    run.text = run.text.replace(placeholder, replacement)
                    # Применяем сохраненный стиль
                    run.style = style
                    # Добавляем подчеркивание
                    run.font.underline = True
                    # Настройка цвета подчеркивания (необязательно)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет


def replace_placeholder_in_table(doc, placeholder, replacement):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                # Сохраняем стиль текста
                                style = run.style
                                # Заменяем текст
                                run.text = run.text.replace(placeholder, replacement)
                                # Применяем сохраненный стиль
                                run.style = style
                                # Добавляем подчеркивание
                                run.font.underline = True
                                # Настройка цвета подчеркивания (необязательно)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет


def add_date_to_end(doc):
    for paragraph in doc.paragraphs:
        if "Дата формирования претензии" in paragraph.text:
            # Устанавливаем локаль на русский
            locale.setlocale(locale.LC_TIME, "ru_RU.UTF-8")
            current_date = datetime.now().strftime("%d %B %Y")
            paragraph.text = paragraph.text.replace("___", current_date)
            break


docx_file_path = "требование WB.docx"
new_docx_file_path = "новый_требование WB.docx"

doc = Document(docx_file_path)

fields = [
    {"placeholder": "поле1", "replacement_text": "МБА ПЕРФЮМЕС МАНУФАКТУРИНГ ЛЛК"},
    {"placeholder": "поле2", "replacement_text": '''Объединенные Арабские Эмираты
Стор N 4, Малик Мубарак Рашид Мубарак, 
Рас Ал Кхор Индастриал Аре 1, П.О. Бокс 390789, 
Дубай, Объединенные Арабские Эмираты (AE)'''},
    {"placeholder": "поле3", "replacement_text": "zaripovtf@gmail.com"},
    {"placeholder": "поле4", "replacement_text": "+7 (985) 411-27-14"},
    {"placeholder": "поле5", "replacement_text": "адвоката Зарипова Т.Ф."},
    {"placeholder": "поле8", "replacement_text": "адвокат Зарипов Т.Ф."},
    {"placeholder": "поле6", "replacement_text": "945815"},
    {"placeholder": "поле7", "replacement_text": "03 класса МКТУ, включая: «амбра [парфюмерия]; ароматизаторы [эфирные масла]; вода туалетная; духи; изделия парфюмерные»"}
]

for field in fields:
    replace_placeholder_in_table(doc, field["placeholder"], field["replacement_text"])
    replace_placeholder_in_doc(doc, field["placeholder"], field["replacement_text"])

add_date_to_end(doc)

doc.save(new_docx_file_path)
